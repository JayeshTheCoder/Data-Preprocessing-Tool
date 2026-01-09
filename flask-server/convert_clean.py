import os
import sys
import markdown
import requests
import io
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from concurrent.futures import ThreadPoolExecutor, as_completed

def handle_element(element, doc, md_file_dir):
    """
    Recursively parses a BeautifulSoup element and adds its content to the docx Document.
    """
    if element.name is None: # NavigableString
        return

    # --- Block Level Elements ---
    if element.name.startswith('h') and element.name[1:].isdigit():
        level = int(element.name[1:])
        doc.add_heading(element.get_text(strip=True), level=level)

    elif element.name == 'p':
        p = doc.add_paragraph()
        for content in element.contents:
            if content.name is None: # Plain text
                p.add_run(str(content))
            elif content.name in ['strong', 'b']:
                p.add_run(content.get_text(strip=True)).bold = True
            elif content.name in ['em', 'i']:
                p.add_run(content.get_text(strip=True)).italic = True
            elif content.name in ['del', 's']:
                p.add_run(content.get_text(strip=True)).strike = True
            elif content.name == 'code':
                run = p.add_run(content.get_text(strip=True))
                run.font.name = 'Courier New'
            elif content.name == 'a':
                text = content.get_text(strip=True)
                href = content.get('href', '')
                p.add_run(f"{text} ({href})")

    elif element.name == 'ul':
        for li in element.find_all('li', recursive=False):
            item_text = li.get_text(strip=True)
            doc.add_paragraph(item_text, style='List Bullet')

    elif element.name == 'ol':
        for li in element.find_all('li', recursive=False):
            item_text = li.get_text(strip=True)
            doc.add_paragraph(item_text, style='List Number')
            
    elif element.name == 'blockquote':
        doc.add_paragraph(element.get_text(strip=True), style='Intense Quote')

    elif element.name == 'pre' and element.code:
        code_text = element.code.get_text()
        p = doc.add_paragraph(style='No Spacing')
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        
    elif element.name == 'img':
        src = element.get('src')
        if not src: return
        try:
            if src.startswith(('http://', 'https://')):
                response = requests.get(src, stream=True)
                response.raise_for_status()
                image_stream = io.BytesIO(response.content)
                doc.add_picture(image_stream, width=Inches(5.0))
            else:
                image_path = os.path.join(md_file_dir, src)
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(5.0))
        except Exception as e:
            print(f"Warning: Could not add image {src}. Reason: {e}")

    elif element.name == 'table':
        rows = element.find_all('tr')
        if not rows: return
        
        header_cells = rows[0].find_all(['th', 'td'])
        if not header_cells: return
        
        num_cols = len(header_cells)
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        
        for i, cell in enumerate(header_cells):
            table.cell(0, i).text = cell.get_text(strip=True)
        
        for row_data in rows[1:]:
            row_cells_data = row_data.find_all('td')
            new_row = table.add_row().cells
            for i, cell in enumerate(row_cells_data):
                if i < num_cols:
                    new_row[i].text = cell.get_text(strip=True)

    elif element.name == 'hr':
        p = doc.add_paragraph()
        p.add_run('_________________')

def apply_indentation_by_numbering(doc):
    """
    Iterates through all paragraphs in the document.
    Detects numbering patterns (e.g., A.1, 1.1.2) and applies indentation
    based on the hierarchy depth (number of dots).
    
    Logic:
    - 'A'       -> 0 dots -> Level 0
    - 'A.1'     -> 1 dot  -> Level 1 (Indent 0.5")
    - 'A.1.1'   -> 2 dots -> Level 2 (Indent 1.0")
    """
    # Regex to catch numbering at start of line (e.g., "A.1", "1.1.2", "A.0")
    # Captures groups like: "A", "A.1", "1.1.1"
    numbering_pattern = re.compile(r"^\s*([A-Za-z0-9]+(?:\.[A-Za-z0-9]+)*)")

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        match = numbering_pattern.match(text)
        if match:
            numbering_str = match.group(1)
            
            # Calculate level based on number of dots
            # Example: "A" (0 dots), "A.1" (1 dot), "A.1.1" (2 dots)
            indent_level = numbering_str.count('.')
            
            if indent_level > 0:
                paragraph.paragraph_format.left_indent = Inches(0.5 * indent_level)

def remove_header_from_markdown(md_content):
    """
    Removes content before '## 💼 Financial Controller Commentary' marker.
    """
    marker = "## 💼 Financial Controller Commentary"
    if marker in md_content:
        return marker + md_content.split(marker)[1]
    else:
        lines = md_content.split('\n')
        if len(lines) > 8:
            return '\n'.join(lines[8:])
        else:
            return md_content

def convert_md_to_docx(md_file_path):
    if not os.path.exists(md_file_path):
        return f"Error: File not found - {md_file_path}"

    file_name, _ = os.path.splitext(md_file_path)
    docx_file_path = f"{file_name}.docx"
    md_file_dir = os.path.dirname(md_file_path)

    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        cleaned_md_content = remove_header_from_markdown(md_content)
        
        # Convert to HTML
        html = markdown.markdown(cleaned_md_content, extensions=['tables', 'fenced_code', 'sane_lists'])
        soup = BeautifulSoup(html, "html.parser")
        doc = Document()
        
        # Build Document
        for element in soup.find_all(True, recursive=False):
            handle_element(element, doc, md_file_dir)
        
        # --- NEW STEP: Post-process indentation based on numbering ---
        apply_indentation_by_numbering(doc)
        
        doc.save(docx_file_path)
        return f"✓ Successfully converted {os.path.basename(md_file_path)}"
    except Exception as e:
        return f"✗ Error converting {os.path.basename(md_file_path)}: {e}"

def main(folder_path):
    if not os.path.isdir(folder_path):
        print(f"Error: The specified folder does not exist: {folder_path}")
        return

    md_files = [os.path.join(folder_path, f) for f in os.listdir(folder_path) if f.endswith('.md')]

    if not md_files:
        print(f"No Markdown files (.md) found in {folder_path}")
        return

    print(f"Found {len(md_files)} Markdown file(s) to convert in '{os.path.basename(folder_path)}'")
    print("Removing header sections, detecting numbering patterns, and converting...\n")

    with ThreadPoolExecutor(max_workers=os.cpu_count()) as executor:
        future_to_file = {executor.submit(convert_md_to_docx, md_file): md_file for md_file in md_files}

        for future in as_completed(future_to_file):
            result = future.result()
            print(result)

    print("\n✓ Conversion process finished!")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python convert_clean.py <path_to_folder>")
        sys.exit(1)

    input_folder = sys.argv[1]
    main(input_folder)