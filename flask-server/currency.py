import re
import os
import glob
from docx import Document

# --- CONFIGURATION ---
# SET YOUR FOLDER PATH HERE
TARGET_FOLDER_PATH = r'C:\Users\singh-128\Downloads\docx'  
# --- END CONFIGURATION ---


# --- Currency Mapping ---
CURRENCY_MAPPING = {
    'UK': '£',
    'US': '$',
    'EU': '€',
    'JP': '¥',      # Japan
    'IN': '₹',      # India
    'AU': 'A$',     # Australia
    'CA': 'C$',     # Canada
    'CH': 'CHF',    # Switzerland
    'CN': '¥',      # China
    'TH': '฿',      # Thailand based on your filename 'TH01'
}

def get_unit_from_filename(filename):
    """
    Extracts the 2-letter country code from the filename.
    Example: '..._JP01_...docx' -> 'JP'
    """
    # This regex looks for a pattern like '_XX##_' where XX is two capital letters.
    match = re.search(r'_([A-Z]{2})\d+', filename)
    if match:
        return match.group(1)  # Returns 'JP'
    
    print(f"Warning: Could not find country code in '{filename}'.")
    return None

def convert_currency_in_file(filename):
    """
    Reads a .docx file, replaces dollar amounts with the correct currency symbol
    based on the unit name in the filename, and overwrites the file.
    """
    # 1. Determine the unit and currency symbol
    # os.path.basename gets just the filename (e.g., 'OE_Data_Processed_JP01_5031_1025.docx')
    # from the full path
    file_basename = os.path.basename(filename)
    unit_name = get_unit_from_filename(file_basename)
    
    if not unit_name:
        print(f"Error: Could not extract a valid unit name from '{file_basename}'.")
        print("Expected format like '..._XX##_...docx'.")
        return

    currency_symbol = CURRENCY_MAPPING.get(unit_name)
    if not currency_symbol:
        print(f"Error: No currency mapping found for unit '{unit_name}'.")
        print(f"Please add '{unit_name}' to the CURRENCY_MAPPING dictionary.")
        return

    if currency_symbol == '$':
        print(f"Unit '{unit_name}' uses '$'. No changes needed for '{file_basename}'.")
        return

    # 2. Read the .docx file
    try:
        document = Document(filename)
    except Exception as e:
        print(f"Error: Could not open '{filename}'. Is it a valid .docx file? Error: {e}")
        return

    # 3. Define the pattern and replacement
    pattern = r'\$([\d-])'
    replacement = f'{currency_symbol}\\1'
    
    changes_made = False

    try:
        # 4. Perform replacement in all paragraphs
        for p in document.paragraphs:
            if re.search(pattern, p.text):  # <-- FIXED THIS LINE
                p.text = re.sub(pattern, replacement, p.text)
                changes_made = True
        
        # 5. Perform replacement in all tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if re.search(pattern, p.text):  # <-- FIXED THIS LINE
                            p.text = re.sub(pattern, replacement, p.text)
                            changes_made = True

        # 6. Write the updated content back to the file
        if changes_made:
            document.save(filename)
            print(f"Successfully converted currency in '{file_basename}' to {currency_symbol}.")
        else:
            print(f"No currency symbols needing conversion were found in '{file_basename}'.")

    except Exception as e:
        print(f"An error occurred while processing or saving '{filename}': {e}")


def main():
    """
    Main function to find and process all .docx files in a fixed folder.
    """
    search_pattern = os.path.join(TARGET_FOLDER_PATH, '*.docx')
    docx_files = glob.glob(search_pattern)
    
    if not docx_files:
        print(f"No .docx files found in the folder: {TARGET_FOLDER_PATH}")
        return

    print(f"Found {len(docx_files)} .docx file(s) to process in '{TARGET_FOLDER_PATH}'...")

    for filepath in docx_files:
        convert_currency_in_file(filepath)
    
    print("\nProcessing complete.")


if __name__ == '__main__':
    main()