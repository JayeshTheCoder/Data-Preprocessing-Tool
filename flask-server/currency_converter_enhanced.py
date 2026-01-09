"""
Enhanced Currency Symbol Converter (Dynamic)
"""

import re
import os
import sys
import glob
from docx import Document

# --- Currency Mapping ---
CURRENCY_MAPPING = {
    'US': '$', 'CA': 'C$', 'BR': 'R$', 'MX': 'MXN$', 'AR': 'ARS$', 'CL': 'CLP$',
    'UK': '£', 'CH': 'CHF', 'AT': '€', 'BE': '€', 'DE': '€', 'IT': '€', 'ES': '€',
    'FR': '€', 'NL': '€', 'IE': '€', 'FI': '€', 'PT': '€', 'GR': '€', 'EU': '€',
    'JP': '¥', 'CN': 'CN¥', 'KR': '₩', 'SG': 'S$', 'IN': '₹', 'AU': 'A$',
    'ZA': 'R', 'TR': '₺', 'AE': 'AED', 'SA': 'SR', 'TH': '฿', 'MY': 'RM',
    # Add other mappings as needed
}

def get_unit_from_filename(filename):
    """Extracts the 2-letter country code from the filename."""
    if 'CNY' in filename or '(Cons CNY)' in filename or 'China' in filename:
        return 'CN'
    
    special_cases = {
        'MT-B': 'BE', 'PI-D': 'DE', 'VOL-EU': 'EU', 'AM-HUB': 'US',
        'AP-HUB': 'SG', 'SLXR': 'UK', 'PIUS PO': 'US', 'PIUSMO': 'US',
        'MTTJ': 'US', 'THOR': 'US'
    }
    
    for key, country_code in special_cases.items():
        if key in filename: return country_code
    
    match = re.search(r'_([A-Z]{2})\d+', filename)
    if match: return match.group(1)
    
    match = re.search(r'-([A-Z]{2})_', filename)
    if match: return match.group(1)
    
    return None

def convert_currency_in_runs(paragraph, pattern, replacement):
    replacements = 0
    for run in paragraph.runs:
        if re.search(pattern, run.text):
            old_text = run.text
            run.text = re.sub(pattern, replacement, run.text)
            replacements += old_text.count('$')
    return replacements

def convert_currency_in_file(filename, dry_run=False):
    file_basename = os.path.basename(filename)
    unit_name = get_unit_from_filename(file_basename)
    
    result = {'filename': file_basename, 'unit': unit_name, 'status': 'unknown', 'message': '', 'replacements': 0}
    
    if not unit_name:
        result['status'] = 'skipped'
        result['message'] = 'No country code found'
        return result
    
    currency_symbol = CURRENCY_MAPPING.get(unit_name, '$')
    if currency_symbol == '$':
        result['status'] = 'skipped'
        result['message'] = 'Already uses $'
        return result
    
    try:
        document = Document(filename)
        pattern = r'\$(?=[\d-])'
        replacements_made = 0
        
        if dry_run:
            # Simple count estimation for dry run
            count = sum(p.text.count('$') for p in document.paragraphs)
            result['status'] = 'dry_run'
            result['message'] = f'Would convert ~{count} symbols to {currency_symbol}'
            return result

        for p in document.paragraphs:
            if '$' in p.text:
                replacements_made += convert_currency_in_runs(p, pattern, currency_symbol)
        
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if '$' in p.text:
                            replacements_made += convert_currency_in_runs(p, pattern, currency_symbol)
        
        document.save(filename)
        result['status'] = 'success'
        result['message'] = f'Converted {replacements_made} $ to {currency_symbol}'
        result['replacements'] = replacements_made
        
    except Exception as e:
        result['status'] = 'error'
        result['message'] = f'Error: {e}'
    
    return result

def process_input(input_path, dry_run=False):
    """
    Process a specific file or folder for currency conversion.
    """
    files_to_process = []
    
    if os.path.isfile(input_path):
        if input_path.endswith('.docx'):
            files_to_process.append(input_path)
    elif os.path.isdir(input_path):
        for root, dirs, files in os.walk(input_path):
            for file in files:
                if file.endswith('.docx'):
                    files_to_process.append(os.path.join(root, file))
    
    print(f"Found {len(files_to_process)} DOCX file(s) for currency check.")
    
    all_results = []
    for filepath in files_to_process:
        result = convert_currency_in_file(filepath, dry_run=dry_run)
        icon = '✓' if result['status'] == 'success' else '•'
        print(f"{icon} {os.path.basename(filepath)}: {result['message']}")
        all_results.append(result)
        
    return all_results

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python currency_converter_enhanced.py <input_path> [--dry-run]")
        sys.exit(1)
        
    path = sys.argv[1]
    dry_run = '--dry-run' in sys.argv
    process_input(path, dry_run=dry_run)