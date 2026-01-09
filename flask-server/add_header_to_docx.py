"""
Add Header to DOCX Files (Dynamic - Fixed)
"""

import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Configuration for Header
CURRENT_MONTH = "October"
CURRENT_YEAR = "2025"

def add_header_to_document(docx_path, dry_run=False):
    filename = os.path.basename(docx_path)
    result = {'filename': filename, 'status': 'unknown', 'message': ''}
    
    if dry_run:
        result['status'] = 'dry_run'
        result['message'] = f'Would add header to {filename}'
        return result
    
    try:
        doc = Document(docx_path)
        
        # Check for existing header
        if doc.paragraphs and 'Financial Controller Commentary' in doc.paragraphs[0].text:
            result['status'] = 'skipped'
            result['message'] = 'Header already exists'
            return result
        
        # Insert Header
        if doc.paragraphs:
            new_para = doc.paragraphs[0].insert_paragraph_before()
        else:
            new_para = doc.add_paragraph()

        header_text = f"💼 Financial Controller Commentary - {CURRENT_MONTH} {CURRENT_YEAR}"
        run = new_para.add_run(header_text)
        
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        if len(doc.paragraphs) > 1:
            doc.paragraphs[1].insert_paragraph_before()
        
        doc.save(docx_path)
        
        result['status'] = 'success'
        result['message'] = 'Header added successfully'
        
    except Exception as e:
        result['status'] = 'error'
        result['message'] = f'Error: {e}'
    
    return result

def process_input(input_path, dry_run=False):
    files_to_process = []
    
    if os.path.isfile(input_path):
        if input_path.endswith('.docx'):
            files_to_process.append(input_path)
    elif os.path.isdir(input_path):
        for root, dirs, files in os.walk(input_path):
            for file in files:
                # [FIX] Removed check for 'C1_Analysis' so ALL docx files get processed
                if file.endswith('.docx'):
                    files_to_process.append(os.path.join(root, file))
    
    print(f"Found {len(files_to_process)} DOCX file(s) for headers.")
    
    results = []
    for filepath in files_to_process:
        result = add_header_to_document(filepath, dry_run=dry_run)
        print(f"{'✓' if result['status']=='success' else '•'} {os.path.basename(filepath)}: {result['message']}")
        results.append(result)
        
    return results

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python add_header_to_docx.py <input_path> [--dry-run]")
        sys.exit(1)
        
    path = sys.argv[1]
    dry_run = '--dry-run' in sys.argv
    process_input(path, dry_run=dry_run)