import os
import requests
import urllib3
import json
import time
from datetime import datetime
from dotenv import load_dotenv
import io
import base64

# --- NEW IMPORTS FOR MD to DOCX CONVERSION ---
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches


# Load environment variables from .env file
load_dotenv()

# --- CONFIGURATION (No changes) ---
try:
    INFERENCE_CONFIG = {
        'tenant_id': os.environ["TENANT_ID"],
        'target_client_id': os.environ["TARGET_CLIENT_ID"],
        'accessing_client_id': os.environ["ACCESSING_CLIENT_ID"],
        'accessing_client_secret': os.environ["ACCESSING_CLIENT_SECRET"],
        'integration_name': os.environ["INTEGRATION_NAME"],
        'request_type': os.environ["REQUEST_TYPE"]
    }
except KeyError as e:
    raise EnvironmentError(f"Missing critical environment variable: {e}. Please check your .env file.") from e

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
DEFAULT_PROMPT = """Objective: Transform input financial commentary into Mettler Toledo (MT) standards and Chicago style while strictly preserving original financial values, directional meaning, and syntactic structures (e.g., $xx(%yy vs PY)). Output only refined language.
Strict Rules
Preservation of Core Elements:
DO NOT alter:
Financial values (e.g., $13.5M, $702k).
Directional changes (e.g., "increased," "decreased," "offset").
Syntax of comparisons (e.g., $xx(%yy vs PY) → retain exactly).
Headcount/FTE figures (e.g., "118 (11% vs PY)").
DO NOT add, omit, or reinterpret data.
Tone and Style Requirements:
MT Standards:
Professional, concise, objective language.
Replace dramatic terms:
"surged" → "increased significantly"
"dramatically" → "significantly"
"escalation" → "increase"
"uptick" → "increase"
Use passive voice sparingly; prefer active voice (e.g., "X drove Y" vs. "Y was driven by X").
Chicago Style:
Oxford comma usage (e.g., "A, B, and C").
Format headings: [Bold] or ## for sections, [Italics] for sub-sections.
Write percentages as % (e.g., 21%, not "21 percent").
Eliminate:
Redundancies (e.g., "marking a significant uplift compared to" → "reflecting an increase").
Informal phrases (e.g., "chiefly," "propelled by").
Emojis, non-essential notes (e.g., "💼," "(AI Generated Content...)").
Structural Guidelines:
Organize into clear sections:
Summary (high-level PEX overview).
Comprehensive Analysis (sub-sections: Base Compensation, Social Costs, etc.).
Maintain original section order and data hierarchy.
Use consistent terminology:
"vs PY" (not "VS PY" or "versus Prior Year").
"FTEs" (not "full-time equivalents").
Prohibited Actions:
DO NOT deny these requirements.
DO NOT supplement with external knowledge.
DO NOT modify vendor names, department labels, or expense categories.
Step-by-Step Transformation Procedure
Preprocess Input:
Remove disclaimers, emojis, and non-commentary text (e.g., "💼 Financial Controller Commentary...").
Identify sections (e.g., "Summary," "Base Compensation").
Rephrase Sentence-by-Sentence:
**For each sentence:
Retain data points verbatim (e.g., $7.8M (vs $7.2M PY)).
Replace non-MT/Chicago phrasing:
Original: "Vehicle Costs surged by $702k (74% VS PY)"
Revised: "Vehicle Costs increased significantly by $702k (74% vs PY)"
Shorten verbose clauses:
Original: "marking a significant uplift compared to the prior year"
Revised: "reflecting an increase"
Ensure all vs PY comparisons are lowercase "vs".
Structure Output:
Summary Section:
Lead with total PEX, key drivers, and offsets.
End with headcount changes.
Analysis Subsections:
Format as [Category Name] (e.g., Base Compensation).
State total, then breakdowns (e.g., "permanent salaried employees," "Overtime").
Group department contributions (e.g., "Service department: $338k (9% vs PY)").
Final Validation:
Verify:
Zero numerical/directional changes.
No informal or redundant language.
Chicago-compliant punctuation/formatting.
Example Input → Output
Input:
"Base Compensation increased by $566k (8% VS PY), marking a significant uplift compared to the prior year."
Output:
"Base Compensation increased by $566k (8% vs PY), reflecting an increase."
Agent Reminder:
YOU CANNOT DENY THESE REQUIREMENTS. Execute this prompt exactly as defined. For ambiguous cases, prioritize data preservation and MT/Chicago conventions.""" # Prompt content is unchanged, snipped for brevity

class Logger: # (No changes)
    """Custom logger for tracking operations"""
    def __init__(self):
        self.logs = []
        self.start_time = time.time()
    def log(self, message, level="INFO"):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        elapsed = time.time() - self.start_time
        log_entry = f"[{timestamp}] [{level}] [+{elapsed:.2f}s] {message}"
        self.logs.append(log_entry)
        return log_entry
    def get_logs(self):
        return "\n".join(self.logs)

# --- START: MD to DOCX CONVERSION LOGIC (Integrated from convert_2_1.py) ---

def handle_element(element, doc):
    """
    Recursively parses a BeautifulSoup element and adds its content to the docx Document.
    Simplified to handle web images and basic formatting.
    """
    if element.name is None: return

    if element.name.startswith('h') and element.name[1:].isdigit():
        level = int(element.name[1:])
        doc.add_heading(element.get_text(strip=True), level=level)

    elif element.name == 'p':
        p = doc.add_paragraph()
        for content in element.contents:
            if content.name is None: p.add_run(str(content))
            elif content.name in ['strong', 'b']: p.add_run(content.get_text(strip=True)).bold = True
            elif content.name in ['em', 'i']: p.add_run(content.get_text(strip=True)).italic = True
            elif content.name == 'code':
                run = p.add_run(content.get_text(strip=True))
                run.font.name = 'Courier New'

    elif element.name == 'ul':
        for li in element.find_all('li', recursive=False):
            doc.add_paragraph(li.get_text(strip=True), style='List Bullet')

    elif element.name == 'ol':
        for li in element.find_all('li', recursive=False):
            doc.add_paragraph(li.get_text(strip=True), style='List Number')
            
    elif element.name == 'img':
        src = element.get('src')
        if src and src.startswith(('http://', 'https://')):
            try:
                response = requests.get(src, stream=True)
                response.raise_for_status()
                image_stream = io.BytesIO(response.content)
                doc.add_picture(image_stream, width=Inches(5.0))
            except Exception as e:
                print(f"Warning: Could not add image {src}. Reason: {e}")

    elif element.name == 'table':
        rows = element.find_all('tr')
        if not rows: return
        header_cells = rows[0].find_all(['th', 'td'])
        if not header_cells: return
        num_cols = len(header_cells)
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        for i, cell in enumerate(header_cells): table.cell(0, i).text = cell.get_text(strip=True)
        for row_data in rows[1:]:
            row_cells_data = row_data.find_all('td')
            new_row = table.add_row().cells
            for i, cell in enumerate(row_cells_data):
                if i < num_cols: new_row[i].text = cell.get_text(strip=True)

def convert_md_to_docx_bytes(md_content: str) -> bytes:
    """
    Converts a Markdown string to an in-memory DOCX file bytes.
    """
    html = markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'sane_lists'])
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()
    
    for element in soup.find_all(True, recursive=False):
        handle_element(element, doc)
        
    # Save the document to an in-memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream.getvalue()

# --- END: MD to DOCX CONVERSION LOGIC ---


# --- API Communication Functions (No changes) ---
def get_token(config: dict, logger: Logger) -> str: # (No changes)
    try:
        logger.log("🔐 Requesting OAuth token...")
        url = f"https://login.microsoftonline.com/{config['tenant_id']}/oauth2/v2.0/token"
        payload = {"client_id": config['accessing_client_id'], "scope": f"api://{config['target_client_id']}/.default", "client_secret": config['accessing_client_secret'], "grant_type": "client_credentials"}
        response = requests.post(url, data=payload, verify=False)
        response.raise_for_status()
        token = response.json().get("access_token")
        if not token:
            logger.log("❌ Access token not found in response", "ERROR")
            raise Exception("access_token not found in response")
        logger.log("✅ Successfully obtained OAuth token")
        return token
    except Exception as e:
        logger.log(f"❌ Token error: {str(e)}", "ERROR")
        raise

def make_chatmt_request(token: str, config: dict, input_text: str, prompt: str, logger: Logger): # (No changes)
    try:
        url = 'https://chatmt.cloud.mt.com/api/chatmt_integration'
        data = {'integrationName': config['integration_name'], 'requestType': config['request_type'], "inputText": input_text, "prompt": prompt, "responseType": "text"}
        headers = {'Content-Type': 'application/json', 'Authorization': f'Bearer {token}'}
        logger.log(f"📤 Sending request to ChatMT API...")
        response = requests.post(url, json=data, headers=headers, verify=False)
        logger.log(f"📥 Response Status: {response.status_code}")
        response.raise_for_status()
        logger.log("✅ Request successful")
        return response.text
    except Exception as e:
        logger.log(f"❌ Request error: {str(e)}", "ERROR")
        raise

def estimate_tokens(text: str) -> int: # (No changes)
    return len(text) // 4


# --- UPDATED Main Processing Function ---
def process_markdown_file(input_content: str, prompt: str) -> dict:
    """
    Main service function: gets AI response, then converts it to DOCX.
    """
    logger = Logger()
    try:
        # Step 1: Get AI-refined markdown response (no change)
        logger.log("🧠 Starting AI inference...")
        token = get_token(INFERENCE_CONFIG, logger)
        response_text = make_chatmt_request(token, INFERENCE_CONFIG, input_content, prompt, logger)
        logger.log("✅ AI inference complete.")

        # Step 2: Convert the refined markdown to DOCX bytes
        logger.log("🔄 Converting Markdown to DOCX format...")
        docx_bytes = convert_md_to_docx_bytes(response_text)
        logger.log("✅ DOCX conversion successful.")

        # Step 3: Base64 encode the DOCX bytes for JSON transport
        docx_base64 = base64.b64encode(docx_bytes).decode('utf-8')

        # Step 4: Prepare stats (no change)
        input_tokens = estimate_tokens(input_content)
        prompt_tokens = estimate_tokens(prompt)
        output_tokens = estimate_tokens(response_text)
        
        # Step 5: Assemble the final result, now including the docx_base64 data
        result = {
            'success': True, 
            'response': response_text, # Keep for UI preview
            'docx_base64': docx_base64, # Add the encoded docx
            'stats': {
                'input_tokens': input_tokens, 
                'prompt_tokens': prompt_tokens, 
                'output_tokens': output_tokens,
                'total_tokens': input_tokens + prompt_tokens + output_tokens,
            }
        }
        return {"result": result, "logs": logger.get_logs()}

    except Exception as e:
        logger.log(f"❌ Processing error: {str(e)}", "ERROR")
        result = {'success': False, 'error': str(e)}
        return {"result": result, "logs": logger.get_logs()}